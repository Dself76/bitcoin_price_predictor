import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import docx
from textblob import TextBlob
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import classification_report#first time using this so look at this first

class SentimentAnalyzer:
    def __init__(self, docx_file, csv_file):
        self.docx_file = docx_file
        self.csv_file = csv_file
        self.df = None
        plt.style.use('ggplot')

    def extract_text_from_docx(self):
        # Extract text from the DOCX file and save to CSV
        doc = docx.Document(self.docx_file)
        data = []
        for word in doc.paragraphs:
            if word.text.strip():
                data.append([word.text])
        self.df = pd.DataFrame(data, columns=['text'])
        self.df.to_csv(self.csv_file, index=False)
        print("All document text has been saved to CSV.")

    def load_csv(self):
        # Load the CSV file and make sure text is in the correct format
        self.df = pd.read_csv(self.csv_file)
        self.df['text'] = self.df['text'].astype(str)

    def perform_sentiment_analysis(self):
        # Perform sentiment analysis using TextBlob
        self.df['sentiment'] = self.df['text'].apply(lambda x: TextBlob(x).sentiment.polarity)

    def visualize_sentiment_distribution(self):
        # Visualize the distribution of sentiment scores
        plt.figure(figsize=(10, 6))
        sns.histplot(self.df['sentiment'], kde=True)
        plt.title('Sentiment Distribution')
        plt.xlabel('Sentiment Polarity')
        plt.ylabel('Frequency')
        plt.show()

    def calculate_average_sentiment(self):
        # Calculate and display the average sentiment
        average_sentiment = self.df['sentiment'].mean()
        print(f"Average sentiment: {average_sentiment:.2f}")

    def categorize_sentiments(self):
        # Categorize sentiments into Positive, Negative, and Neutral, not sure How i want to handle this yet
        def categorize_sentiment(polarity):
            if polarity > 0.05:
                return 'Positive'
            elif polarity < -0.05:
                return 'Negative'
            else:
                return 'Neutral'
        self.df['sentiment_category'] = self.df['sentiment'].apply(categorize_sentiment)

    def visualize_sentiment_categories(self):
        # going to graph the distribution of sentiment categories so I see
        sentiment_counts = self.df['sentiment_category'].value_counts()
        plt.figure(figsize=(8, 6))
        sentiment_counts.plot(kind='bar')
        plt.title('Sentiment Category Distribution')
        plt.xlabel('Sentiment Category')
        plt.ylabel('Count')
        plt.show()

    def calculate_sentiment_volatility(self):
        # Calculate sentiment volatility using a 7-day rolling window
        self.df['sentiment_volatility'] = self.df['sentiment'].rolling(window=7).std()

    def visualize_sentiment_volatility(self):
        # Visualize sentiment volatility over time..I might have to change this depending on the coin or stock,
        plt.figure(figsize=(10, 6))
        plt.plot(self.df.index, self.df['sentiment_volatility'])
        plt.title('Sentiment Volatility Over Time')
        plt.xlabel('Date')
        plt.ylabel('Sentiment Volatility')
        plt.show()

    def prepare_data_for_ml(self):
        # Prepare data for machine learning by creating sentiment signals and lagged features
        def get_sentiment_signal(sentiment):
            if sentiment > 0.1:
                return 1  # Buy
            elif sentiment < -0.1:
                return -1  # Sell
            else:
                return 0  # Hold
        self.df['sentiment_signal'] = self.df['sentiment'].apply(get_sentiment_signal)
        for i in range(1, 8):
            self.df[f'sentiment_lag_{i}'] = self.df['sentiment'].shift(i)

    def train_and_evaluate_model(self):
        # Train and evaluate a machine learning model
        features = ['sentiment', 'sentiment_volatility'] + [f'sentiment_lag_{i}' for i in range(1, 8)]
        X = self.df[features].dropna()
        y = self.df['sentiment_signal'].loc[X.index]
        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
        scaler = StandardScaler()
        X_train_scaled = scaler.fit_transform(X_train)
        X_test_scaled = scaler.transform(X_test)
        model = LogisticRegression(random_state=42)
        model.fit(X_train_scaled, y_train)
        y_pred = model.predict(X_test_scaled)
        print(classification_report(y_test, y_pred))
        joblib.dump(model, 'sentiment_model.joblib')
        joblib.dump(scaler, 'sentiment_scaler.joblib')
        print("Model and scaler saved for future use.")

    def run_analysis(self):
        try:
            print("Starting analysis...")
            self.extract_text_from_docx()
            self.load_csv()
            self.perform_sentiment_analysis()
            self.visualize_sentiment_distribution()
            self.calculate_average_sentiment()
            self.categorize_sentiments()
            self.visualize_sentiment_categories()
            self.calculate_sentiment_volatility()
            self.visualize_sentiment_volatility()
            self.prepare_data_for_ml()
            self.train_and_evaluate_model()
            print("Analysis completed successfully.")
        except Exception as e:
            print(f"An error occurred during analysis: {str(e)}")

if __name__ == "__main__":
    try:
        print("Starting sentiment analysis...")
        analyzer = SentimentAnalyzer('Files\\bitcoinData.docx', 'Files\\data_text.csv')
        analyzer.run_analysis()
        print("Analysis completed successfully.")
    except Exception as e:
        print(f"An error occurred: {str(e)}")

#function below
'''
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import nltk
import docx
from textblob import TextBlob
from collections import Counter
import re
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import classification_report
import joblib

class SentimentAnalyzer:
    def __init__(self, docx_file, csv_file):
        self.docx_file = docx_file
        self.csv_file = csv_file
        self.df = None
        plt.style.use('ggplot')

    def extract_text_from_docx(self):
        # Extract text from the DOCX file and save to CSV
        doc = docx.Document(self.docx_file)
        data = []
        for word in doc.paragraphs:
            if word.text.strip():
                data.append([word.text])
        self.df = pd.DataFrame(data, columns=['text'])
        self.df.to_csv(self.csv_file, index=False)
        print("All document text has been saved to CSV.")

    def load_csv(self):
        # Load the CSV file and ensure text is in the correct format, I will leave this here for now just in case i have problem in main importing this class
        self.df = pd.read_csv(self.csv_file)
        self.df['text'] = self.df['text'].astype(str)

    def perform_sentiment_analysis(self):
        # Perform sentiment analysis using TextBlob
        self.df['sentiment'] = self.df['text'].apply(lambda x: TextBlob(x).sentiment.polarity)

    def visualize_sentiment_distribution(self):
        # Visualize the distribution of sentiment scores, probably take this out later, but i #like it right now so I can see what happening.
        plt.figure(figsize=(10, 6))
        sns.histplot(self.df['sentiment'], kde=True)
        plt.title('Sentiment Distribution')
        plt.xlabel('Sentiment Polarity')
        plt.ylabel('Frequency')
        plt.show()

    def calculate_average_sentiment(self):
        # Calculate and display the average sentiment
        average_sentiment = self.df['sentiment'].mean()
        print(f"Average sentiment: {average_sentiment:.2f}")

    def categorize_sentiments(self):
        # Categorize sentiments into Positive, Negative, and Neutral
        def categorize_sentiment(polarity):
            if polarity > 0.05:
                return 'Positive'
            elif polarity < -0.05:
                return 'Negative'
            else:
                return 'Neutral'
        self.df['sentiment_category'] = self.df['sentiment'].apply(categorize_sentiment)

    def visualize_sentiment_categories(self):
        # Visualize the distribution of sentiment categories
        sentiment_counts = self.df['sentiment_category'].value_counts()
        plt.figure(figsize=(8, 6))
        sentiment_counts.plot(kind='bar')
        plt.title('Sentiment Category Distribution')
        plt.xlabel('Sentiment Category')
        plt.ylabel('Count')
        plt.show()

    def perform_word_frequency_analysis(self):
        # Perform word frequency analysis
        def get_word_freq(text):
            words = re.findall(r'\w+', text.lower())
            return Counter(words)
        self.df['word_freq'] = self.df['text'].apply(get_word_freq)
        all_words = Counter()
        for freq in self.df['word_freq']:
            all_words.update(freq)
        print("Most common words:")
        print(all_words.most_common(10))

    def calculate_sentiment_volatility(self):
        # Calculate sentiment volatility using a 7-day rolling window
        self.df['sentiment_volatility'] = self.df['sentiment'].rolling(window=7).std()

    def visualize_sentiment_volatility(self):
        # Visualize sentiment volatility over time
        plt.figure(figsize=(10, 6))
        plt.plot(self.df.index, self.df['sentiment_volatility'])
        plt.title('Sentiment Volatility Over Time')
        plt.xlabel('Date')
        plt.ylabel('Sentiment Volatility')
        plt.show()

    def prepare_data_for_ml(self):
        # Prepare data for machine learning by creating sentiment signals and lagged features
        def get_sentiment_signal(sentiment):
            if sentiment > 0.1:
                return 1  # Buy
            elif sentiment < -0.1:
                return -1  # Sell
            else:
                return 0  # Hold
        self.df['sentiment_signal'] = self.df['sentiment'].apply(get_sentiment_signal)
        for i in range(1, 8):
            self.df[f'sentiment_lag_{i}'] = self.df['sentiment'].shift(i)

    def train_and_evaluate_model(self):
        # Train and evaluate a machine learning model
        features = ['sentiment', 'sentiment_volatility'] + [f'sentiment_lag_{i}' for i in range(1, 8)]
        X = self.df[features].dropna()
        y = self.df['sentiment_signal'].loc[X.index]
        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
        scaler = StandardScaler()
        X_train_scaled = scaler.fit_transform(X_train)
        X_test_scaled = scaler.transform(X_test)
        model = LogisticRegression(random_state=42)
        model.fit(X_train_scaled, y_train)
        y_pred = model.predict(X_test_scaled)
        print(classification_report(y_test, y_pred))
        joblib.dump(model, 'sentiment_model.joblib')
        joblib.dump(scaler, 'sentiment_scaler.joblib')
        print("Model and scaler saved for future use.")

    def run_analysis(self):
        # Execute the full analysis pipeline
        self.extract_text_from_docx()
        self.load_csv()
        self.perform_sentiment_analysis()
        self.visualize_sentiment_distribution()
        self.calculate_average_sentiment()
        self.categorize_sentiments()
        self.visualize_sentiment_categories()
        self.perform_word_frequency_analysis()
        self.calculate_sentiment_volatility()
        self.visualize_sentiment_volatility()
        self.prepare_data_for_ml()
        self.train_and_evaluate_model()

# Instantiate the SentimentAnalyzer and run the analysis
analyzer = SentimentAnalyzer('Files\\bitcoinData.docx', 'Files\\data_text.csv')
analyzer.run_analysis()

'''
